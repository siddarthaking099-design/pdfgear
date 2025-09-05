from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import PyPDF2
import io
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import pandas as pd
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4, legal
import zipfile
from PIL import Image

app = Flask(__name__)
CORS(app)

class PDFProcessor:
    
    @staticmethod
    def pdf_to_word(pdf_file):
        """Smart PDF to Word conversion with auto-detection"""
        try:
            pdf_file.seek(0)
            
            # Check if PDF is scanned
            is_scanned = PDFProcessor._check_if_scanned_pdf(pdf_file)
            pdf_file.seek(0)
            
            if is_scanned:
                return PDFProcessor._pdf_to_word_ocr(pdf_file)
            else:
                # Try regular extraction first
                try:
                    result = PDFProcessor._pdf_to_word_regular(pdf_file)
                    return result
                except:
                    # Fallback to OCR if regular fails
                    pdf_file.seek(0)
                    return PDFProcessor._pdf_to_word_ocr(pdf_file)
                    
        except Exception as e:
            raise Exception(f"PDF to Word conversion failed: {str(e)}")
    
    @staticmethod
    def _check_if_scanned_pdf(pdf_file):
        """Check if PDF is likely scanned"""
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            # Check first 3 pages
            for i in range(min(3, len(pdf_reader.pages))):
                page = pdf_reader.pages[i]
                text += page.extract_text()
            
            word_count = len(text.split())
            return word_count < 10  # If less than 10 words, likely scanned
            
        except:
            return False  # Default to regular if can't analyze
    
    @staticmethod
    def _pdf_to_word_regular(pdf_file):
        """Regular PDF to Word using pdf2docx"""
        try:
            from pdf2docx import parse
            import tempfile
            
            # Save to temp file for pdf2docx
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                pdf_file.seek(0)
                temp_pdf.write(pdf_file.read())
                temp_pdf_path = temp_pdf.name
            
            # Convert using pdf2docx
            output = io.BytesIO()
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                parse(temp_pdf_path, temp_docx.name)
                
                with open(temp_docx.name, 'rb') as f:
                    output.write(f.read())
            
            # Cleanup
            import os
            os.unlink(temp_pdf_path)
            os.unlink(temp_docx.name)
            
            output.seek(0)
            return output
            
        except ImportError:
            # Fallback to PyMuPDF if pdf2docx not available
            return PDFProcessor._pdf_to_word_fallback(pdf_file)
    
    @staticmethod
    def _pdf_to_word_ocr(pdf_file):
        """OCR-based PDF to Word conversion"""
        try:
            import pdf2image
            import pytesseract
            from docx.shared import Inches, Pt
            
            pdf_bytes = pdf_file.read()
            images = pdf2image.convert_from_bytes(pdf_bytes, dpi=300)
            
            doc = Document()
            
            # Set margins
            for section in doc.sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            total_words = 0
            
            for i, image in enumerate(images):
                if i > 0:
                    doc.add_page_break()
                
                try:
                    # Perform OCR
                    text = pytesseract.image_to_string(image, lang='eng')
                    
                    if text.strip():
                        # Split into paragraphs
                        paragraphs = text.split('\n\n')
                        
                        for para_text in paragraphs:
                            para_text = para_text.strip()
                            if para_text:
                                paragraph = doc.add_paragraph(para_text)
                                paragraph.paragraph_format.space_after = Pt(6)
                                total_words += len(para_text.split())
                    else:
                        doc.add_paragraph(f"[No text detected on page {i+1}]")
                        
                except Exception as e:
                    doc.add_paragraph(f"[OCR Error on page {i+1}: {str(e)}]")
            
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output
            
        except ImportError:
            # Fallback if OCR libraries not available
            return PDFProcessor._pdf_to_word_fallback(pdf_file)
    
    @staticmethod
    def _pdf_to_word_fallback(pdf_file):
        """Enhanced fallback PDF to Word conversion using PyMuPDF"""
        try:
            pdf_bytes = pdf_file.read()
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            doc = Document()
            
            # Set margins
            for section in doc.sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                if page_num > 0:
                    doc.add_page_break()
                
                # Extract text with formatting
                try:
                    blocks = page.get_text("dict")["blocks"]
                    
                    for block in blocks:
                        if "lines" not in block:
                            continue
                        
                        for line in block["lines"]:
                            if not line.get("spans"):
                                continue
                            
                            paragraph = doc.add_paragraph()
                            
                            for span in line["spans"]:
                                text = span.get("text", "")
                                if text.strip():
                                    run = paragraph.add_run(text)
                                    
                                    # Apply basic formatting
                                    font_size = span.get("size", 12)
                                    font_flags = span.get("flags", 0)
                                    
                                    run.font.size = Pt(font_size)
                                    run.font.bold = bool(font_flags & 16)
                                    run.font.italic = bool(font_flags & 2)
                except:
                    # Simple text extraction if structured fails
                    text = page.get_text()
                    if text.strip():
                        lines = text.split('\n')
                        for line in lines:
                            line = line.strip()
                            if line:
                                doc.add_paragraph(line)
                    else:
                        doc.add_paragraph(f"[Page {page_num + 1} - No text found]")
            
            pdf_document.close()
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output
        except Exception as e:
            # Last resort - create document with error message
            doc = Document()
            doc.add_paragraph(f"Error extracting text from PDF: {str(e)}")
            doc.add_paragraph("Please try uploading a different PDF file.")
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output

    @staticmethod
    def pdf_to_excel(pdf_file):
        """Convert PDF to Excel"""
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            data = []
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                lines = text.split('\n')
                for line_num, line in enumerate(lines):
                    if line.strip():
                        data.append({
                            'Page': page_num + 1,
                            'Line': line_num + 1,
                            'Content': line.strip()
                        })
            
            df = pd.DataFrame(data)
            output = io.BytesIO()
            df.to_excel(output, index=False)
            output.seek(0)
            return output
        except Exception as e:
            raise Exception(f"Error converting PDF to Excel: {str(e)}")
    
    @staticmethod
    def scan_pdf_pages(pdf_file):
        """Scan PDF and return page previews"""
        try:
            import base64
            pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
            pages = []
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                pix = page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5), alpha=False)
                img_data = pix.tobytes("png")
                img_base64 = base64.b64encode(img_data).decode('utf-8')
                
                pages.append({
                    'page_num': page_num + 1,
                    'preview': f'data:image/png;base64,{img_base64}'
                })
                pix = None
            
            pdf_document.close()
            return pages
        except Exception as e:
            raise Exception(f"Error scanning PDF: {str(e)}")
    
    @staticmethod
    def convert_selected_pages(pdf_file, selected_pages, image_format='png', quality=2):
        """Convert selected PDF pages to images"""
        try:
            import base64
            pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
            images = []
            
            dpi = 150 if quality == 1 else 200 if quality == 2 else 300
            zoom = dpi / 72.0
            zoom_matrix = fitz.Matrix(zoom, zoom)
            
            for page_num in selected_pages:
                if page_num <= len(pdf_document):
                    page = pdf_document.load_page(page_num - 1)
                    pix = page.get_pixmap(matrix=zoom_matrix, alpha=False)
                    
                    if image_format.lower() == 'jpg':
                        img_data = pix.tobytes("jpeg", jpg_quality=95)
                    else:
                        img_data = pix.tobytes("png")
                    
                    img_base64 = base64.b64encode(img_data).decode('utf-8')
                    images.append({
                        'page_num': page_num,
                        'filename': f'page_{page_num:03d}.{image_format.lower()}',
                        'data': img_base64
                    })
                    pix = None
            
            pdf_document.close()
            return images
        except Exception as e:
            raise Exception(f"Error converting selected pages: {str(e)}")
    
    @staticmethod
    def images_to_pdf(image_files, page_size='letter', fit_mode='fit'):
        """Convert images to PDF"""
        try:
            page_sizes = {'letter': letter, 'a4': A4, 'legal': legal}
            page_size_tuple = page_sizes.get(page_size, letter)
            width, height = page_size_tuple
            
            output = io.BytesIO()
            c = canvas.Canvas(output, pagesize=page_size_tuple)
            
            for image_file in image_files:
                image_file.seek(0)
                img = Image.open(image_file)
                img_width, img_height = img.size
                
                if fit_mode == 'stretch':
                    new_width, new_height = width * 0.9, height * 0.9
                    x, y = width * 0.05, height * 0.05
                elif fit_mode == 'fill':
                    scale = max(width/img_width, height/img_height) * 0.9
                    new_width, new_height = img_width * scale, img_height * scale
                    x = (width - new_width) / 2
                    y = (height - new_height) / 2
                else:
                    scale = min(width/img_width, height/img_height) * 0.8
                    new_width, new_height = img_width * scale, img_height * scale
                    x = (width - new_width) / 2
                    y = (height - new_height) / 2
                
                if img.mode in ('RGBA', 'LA', 'P'):
                    img = img.convert('RGB')
                
                temp_img = io.BytesIO()
                img.save(temp_img, format='JPEG', quality=85)
                temp_img.seek(0)
                
                c.drawImage(temp_img, x, y, new_width, new_height)
                c.showPage()
            
            c.save()
            output.seek(0)
            return output
        except Exception as e:
            raise Exception(f"Error converting images to PDF: {str(e)}")

    @staticmethod
    def merge_pdfs(pdf_files):
        """Merge multiple PDF files"""
        try:
            merger = PyPDF2.PdfMerger()
            for pdf_file in pdf_files:
                pdf_file.seek(0)
                merger.append(pdf_file)
            output = io.BytesIO()
            merger.write(output)
            merger.close()
            output.seek(0)
            return output
        except Exception as e:
            raise Exception(f"Error merging PDFs: {str(e)}")
    
    @staticmethod
    def split_pdf(pdf_file):
        """Split PDF into individual pages with base64 data for direct download"""
        try:
            import base64
            pdf_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            pages = []
            
            for page_num in range(len(pdf_reader.pages)):
                writer = PyPDF2.PdfWriter()
                writer.add_page(pdf_reader.pages[page_num])
                
                page_output = io.BytesIO()
                writer.write(page_output)
                page_output.seek(0)
                
                pdf_base64 = base64.b64encode(page_output.getvalue()).decode('utf-8')
                pages.append({
                    'page_num': page_num + 1,
                    'filename': f'page_{page_num + 1:03d}.pdf',
                    'data': pdf_base64
                })
            
            return pages
        except Exception as e:
            raise Exception(f"Error splitting PDF: {str(e)}")
    
    @staticmethod
    def get_split_page(pdf_file, page_num):
        """Get individual page from PDF"""
        try:
            pdf_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            if page_num < 1 or page_num > len(pdf_reader.pages):
                raise Exception("Invalid page number")
            
            writer = PyPDF2.PdfWriter()
            writer.add_page(pdf_reader.pages[page_num - 1])
            
            output = io.BytesIO()
            writer.write(output)
            output.seek(0)
            return output
        except Exception as e:
            raise Exception(f"Error extracting page: {str(e)}")
    
    @staticmethod
    def convert_selected_pdf_pages(pdf_file, selected_pages):
        """Convert selected PDF pages to individual PDFs with base64 data"""
        try:
            import base64
            pdf_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            pdfs = []
            
            for page_num in selected_pages:
                if page_num <= len(pdf_reader.pages) and page_num > 0:
                    writer = PyPDF2.PdfWriter()
                    writer.add_page(pdf_reader.pages[page_num - 1])
                    
                    page_output = io.BytesIO()
                    writer.write(page_output)
                    page_output.seek(0)
                    
                    pdf_base64 = base64.b64encode(page_output.getvalue()).decode('utf-8')
                    pdfs.append({
                        'page_num': page_num,
                        'filename': f'page_{page_num:03d}.pdf',
                        'data': pdf_base64
                    })
            
            return pdfs
        except Exception as e:
            raise Exception(f"Error converting selected PDF pages: {str(e)}")
    
    @staticmethod
    def rotate_pdf(pdf_file, rotation):
        """Rotate PDF pages"""
        try:
            pdf_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            writer = PyPDF2.PdfWriter()
            
            for page in pdf_reader.pages:
                page.rotate(rotation)
                writer.add_page(page)
            
            output = io.BytesIO()
            writer.write(output)
            output.seek(0)
            return output
        except Exception as e:
            raise Exception(f"Error rotating PDF: {str(e)}")
    
    @staticmethod
    def delete_pages(pdf_file, pages_to_delete):
        """Delete specific pages from PDF"""
        try:
            pdf_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            writer = PyPDF2.PdfWriter()
            
            total_pages = len(pdf_reader.pages)
            pages_to_delete = set(pages_to_delete)
            
            for page_num in range(total_pages):
                if (page_num + 1) not in pages_to_delete:
                    writer.add_page(pdf_reader.pages[page_num])
            
            output = io.BytesIO()
            writer.write(output)
            output.seek(0)
            return output
        except Exception as e:
            raise Exception(f"Error deleting pages: {str(e)}")
    
    @staticmethod
    def compress_pdf(pdf_file):
        """Compress PDF file"""
        try:
            pdf_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            writer = PyPDF2.PdfWriter()
            
            for page in pdf_reader.pages:
                page.compress_content_streams()
                writer.add_page(page)
            
            output = io.BytesIO()
            writer.write(output)
            output.seek(0)
            return output
        except Exception as e:
            raise Exception(f"Error compressing PDF: {str(e)}")
    
    @staticmethod
    def unlock_pdf(pdf_file, password):
        """Remove password protection from PDF"""
        try:
            pdf_file.seek(0)
            
            # Try PyMuPDF first
            try:
                pdf_bytes = pdf_file.read()
                pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
                
                if pdf_document.needs_pass:
                    if not pdf_document.authenticate(password):
                        pdf_document.close()
                        raise Exception("Invalid password")
                
                output = io.BytesIO()
                pdf_document.save(output)
                pdf_document.close()
                output.seek(0)
                return output
                
            except Exception:
                # Fallback to PyPDF2
                pdf_file.seek(0)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                if pdf_reader.is_encrypted:
                    if not pdf_reader.decrypt(password):
                        raise Exception("Invalid password")
                
                writer = PyPDF2.PdfWriter()
                for page in pdf_reader.pages:
                    writer.add_page(page)
                
                output = io.BytesIO()
                writer.write(output)
                output.seek(0)
                return output
                
        except Exception as e:
            raise Exception(f"Error unlocking PDF: {str(e)}")
    
    @staticmethod
    def protect_pdf(pdf_file, password, owner_password=None):
        """Add strong encryption to PDF using PyMuPDF"""
        try:
            pdf_file.seek(0)
            pdf_bytes = pdf_file.read()
            
            # Use PyMuPDF for stronger encryption
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            # Set strong encryption parameters
            encrypt_meth = fitz.PDF_ENCRYPT_AES_256  # AES-256 encryption
            owner_pw = owner_password or password + "_owner"  # Owner password
            user_pw = password  # User password
            
            # Set permissions (restrict printing, copying, etc.)
            permissions = (
                fitz.PDF_PERM_PRINT |  # Allow printing
                fitz.PDF_PERM_COPY |   # Allow copying (can be removed for stricter security)
                fitz.PDF_PERM_ANNOTATE  # Allow annotations
            )
            
            # Apply encryption
            output = io.BytesIO()
            pdf_document.save(
                output,
                encryption=encrypt_meth,
                owner_pw=owner_pw,
                user_pw=user_pw,
                permissions=permissions
            )
            
            pdf_document.close()
            output.seek(0)
            return output
            
        except Exception as e:
            # Fallback to PyPDF2 if PyMuPDF fails
            try:
                pdf_file.seek(0)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                writer = PyPDF2.PdfWriter()
                
                for page in pdf_reader.pages:
                    writer.add_page(page)
                
                # Use PyPDF2 encryption with stronger settings
                writer.encrypt(
                    user_password=password,
                    owner_password=owner_password or password + "_owner",
                    use_128bit=True  # Use 128-bit encryption
                )
                
                output = io.BytesIO()
                writer.write(output)
                output.seek(0)
                return output
            except Exception as fallback_error:
                raise Exception(f"Error protecting PDF: {str(e)} | Fallback error: {str(fallback_error)}")
    
    @staticmethod
    def word_to_pdf(word_file):
        """Convert Word document to PDF"""
        try:
            from docx2pdf import convert
            import tempfile
            import os
            
            # Save uploaded file to temp location
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_word:
                word_file.seek(0)
                temp_word.write(word_file.read())
                temp_word_path = temp_word.name
            
            # Convert to PDF
            temp_pdf_path = temp_word_path.replace('.docx', '.pdf')
            convert(temp_word_path, temp_pdf_path)
            
            # Read PDF and return as BytesIO
            output = io.BytesIO()
            with open(temp_pdf_path, 'rb') as pdf_file:
                output.write(pdf_file.read())
            
            # Cleanup temp files
            os.unlink(temp_word_path)
            os.unlink(temp_pdf_path)
            
            output.seek(0)
            return output
            
        except ImportError:
            raise Exception("docx2pdf library not available. Install with: pip install docx2pdf")
        except Exception as e:
            raise Exception(f"Error converting Word to PDF: {str(e)}")
    
    @staticmethod
    def excel_to_pdf(excel_file):
        """Convert Excel spreadsheet to PDF"""
        try:
            import tempfile
            import os
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter
            
            # Read Excel file
            excel_file.seek(0)
            df = pd.read_excel(excel_file)
            
            # Create PDF
            output = io.BytesIO()
            doc = SimpleDocTemplate(output, pagesize=letter)
            
            # Convert DataFrame to table data
            data = [df.columns.tolist()] + df.values.tolist()
            
            # Create table
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            # Build PDF
            doc.build([table])
            output.seek(0)
            return output
            
        except Exception as e:
            raise Exception(f"Error converting Excel to PDF: {str(e)}")

# API Routes
@app.route('/api/pdf-to-word', methods=['POST'])
def pdf_to_word():
    try:
        file = request.files['file']
        result = PDFProcessor.pdf_to_word(file)
        return send_file(result, as_attachment=True, download_name='converted.docx', 
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/pdf-to-excel', methods=['POST'])
def pdf_to_excel():
    try:
        file = request.files['file']
        result = PDFProcessor.pdf_to_excel(file)
        return send_file(result, as_attachment=True, download_name='converted.xlsx', 
                        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/scan-pdf', methods=['POST'])
def scan_pdf():
    try:
        file = request.files['file']
        pages = PDFProcessor.scan_pdf_pages(file)
        return jsonify({'pages': pages})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/convert-pages', methods=['POST'])
def convert_pages():
    try:
        file = request.files['file']
        selected_pages = list(map(int, request.form.get('pages', '').split(',')))
        image_format = request.form.get('format', 'png')
        quality = int(request.form.get('quality', 2))
        
        images = PDFProcessor.convert_selected_pages(file, selected_pages, image_format, quality)
        return jsonify({'images': images})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/images-to-pdf', methods=['POST'])
def images_to_pdf():
    try:
        files = request.files.getlist('files')
        page_size = request.form.get('page_size', 'letter')
        fit_mode = request.form.get('fit_mode', 'fit')
        result = PDFProcessor.images_to_pdf(files, page_size, fit_mode)
        return send_file(result, as_attachment=True, download_name='images_to_pdf.pdf', 
                        mimetype='application/pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/merge-pdf', methods=['POST'])
def merge_pdf():
    try:
        files = request.files.getlist('files')
        result = PDFProcessor.merge_pdfs(files)
        return send_file(result, as_attachment=True, download_name='merged.pdf', 
                        mimetype='application/pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/split-pdf', methods=['POST'])
def split_pdf():
    try:
        file = request.files['file']
        pages = PDFProcessor.split_pdf(file)
        return jsonify({'pages': pages})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/download-split-page', methods=['POST'])
def download_split_page():
    try:
        file = request.files['file']
        page_num = int(request.form.get('page_num', 1))
        
        result = PDFProcessor.get_split_page(file, page_num)
        return send_file(result, as_attachment=True, 
                        download_name=f'page_{page_num:03d}.pdf', 
                        mimetype='application/pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/convert-selected-pdf-pages', methods=['POST'])
def convert_selected_pdf_pages():
    try:
        file = request.files['file']
        selected_pages = list(map(int, request.form.get('pages', '').split(',')))
        
        if not selected_pages:
            return jsonify({'error': 'No pages selected'}), 400
        
        pdfs = PDFProcessor.convert_selected_pdf_pages(file, selected_pages)
        return jsonify({'pdfs': pdfs})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/download-all-split-pages', methods=['POST'])
def download_all_split_pages():
    try:
        file = request.files['file']
        file.seek(0)
        pdf_reader = PyPDF2.PdfReader(file)
        
        # Create individual PDF files and return as JSON with download links
        pages_info = []
        for page_num in range(len(pdf_reader.pages)):
            pages_info.append({
                'page_num': page_num + 1,
                'filename': f'page_{page_num + 1:03d}.pdf',
                'ready_for_download': True
            })
        
        return jsonify({
            'message': 'PDF split successfully',
            'total_pages': len(pdf_reader.pages),
            'pages': pages_info,
            'download_instruction': 'Use download-split-page endpoint with page_num to download individual pages'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/rotate-pdf', methods=['POST'])
def rotate_pdf():
    try:
        file = request.files['file']
        rotation = int(request.form.get('rotation', 90))
        result = PDFProcessor.rotate_pdf(file, rotation)
        return send_file(result, as_attachment=True, download_name='rotated.pdf', 
                        mimetype='application/pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/delete-pages', methods=['POST'])
def delete_pages():
    try:
        file = request.files['file']
        pages_to_delete = list(map(int, request.form.get('pages', '').split(',')))
        result = PDFProcessor.delete_pages(file, pages_to_delete)
        return send_file(result, as_attachment=True, download_name='pages_deleted.pdf', 
                        mimetype='application/pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/compress-pdf', methods=['POST'])
def compress_pdf():
    try:
        file = request.files['file']
        result = PDFProcessor.compress_pdf(file)
        return send_file(result, as_attachment=True, download_name='compressed.pdf', 
                        mimetype='application/pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/unlock-pdf', methods=['POST'])
def unlock_pdf():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
            
        file = request.files['file']
        password = request.form.get('password', '').strip()
        
        if not password:
            return jsonify({'error': 'Password is required to unlock PDF'}), 400
        
        result = PDFProcessor.unlock_pdf(file, password)
        return send_file(result, as_attachment=True, download_name='unlocked.pdf', 
                        mimetype='application/pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/protect-pdf', methods=['POST'])
def protect_pdf():
    try:
        file = request.files['file']
        password = request.form.get('password', '').strip()
        owner_password = request.form.get('owner_password', '').strip()
        
        if not password:
            password = 'default123'  # Use default password if none provided
        
        result = PDFProcessor.protect_pdf(file, password, owner_password if owner_password else None)
        return send_file(result, as_attachment=True, download_name='protected.pdf', 
                        mimetype='application/pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/word-to-pdf', methods=['POST'])
def word_to_pdf():
    try:
        file = request.files['file']
        result = PDFProcessor.word_to_pdf(file)
        return send_file(result, as_attachment=True, download_name='converted.pdf', 
                        mimetype='application/pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/excel-to-pdf', methods=['POST'])
def excel_to_pdf():
    try:
        file = request.files['file']
        result = PDFProcessor.excel_to_pdf(file)
        return send_file(result, as_attachment=True, download_name='converted.pdf', 
                        mimetype='application/pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/status', methods=['GET'])
def status():
    return jsonify({'status': 'running', 'message': 'PDF Gears Python backend is active'})

if __name__ == '__main__':
    app.run(debug=True, port=5000)